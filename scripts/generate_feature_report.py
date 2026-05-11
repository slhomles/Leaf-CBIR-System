"""Tạo file Word tổng hợp chi tiết quy trình trích xuất đặc trưng của dự án.

Output: reports/feature_extraction_report.docx
Chạy:
    python -X utf8 scripts/generate_feature_report.py
"""

from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "reports", "feature_extraction_report.docx")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, *, size=11, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=10)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_kv_table(doc, rows: list[tuple[str, str]], col1_width=4.5, col2_width=11.5):
    """Thêm bảng 2 cột (key, value)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = False
    for r_idx, (k, v) in enumerate(rows):
        cells = table.rows[r_idx].cells
        cells[0].text = ""
        cells[1].text = ""
        cells[0].width = Cm(col1_width)
        cells[1].width = Cm(col2_width)
        run_k = cells[0].paragraphs[0].add_run(k)
        set_run_font(run_k, size=10, bold=True)
        run_v = cells[1].paragraphs[0].add_run(v)
        set_run_font(run_v, size=10)
    doc.add_paragraph()
    return table


def add_dim_table(doc, headers: list[str], rows: list[list[str]]):
    """Bảng nhiều cột — mỗi cột headers tương ứng."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for c_idx, h in enumerate(headers):
        cell = table.rows[0].cells[c_idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        # nền màu xanh
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "4472C4")
        cell._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=10)
    doc.add_paragraph()
    return table


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

def section_cover(doc):
    title = doc.add_heading("BÁO CÁO TRÍCH XUẤT ĐẶC TRƯNG", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Hệ thống tìm kiếm ảnh lá theo nội dung (Leafsearch CBIR)")
    set_run_font(run, size=14, italic=True, color=(80, 80, 80))

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        f"Tổng vector: 125 chiều  |  4 nhóm đặc trưng riêng biệt\n"
        f"Dataset: 1907 ảnh lá 256×256  |  PostgreSQL + pgvector\n"
        f"Ngày tạo: {date.today().isoformat()}"
    )
    set_run_font(run, size=11)

    doc.add_page_break()


def section_overview(doc):
    add_heading(doc, "1. Tổng quan pipeline", level=1)

    add_para(doc,
        "Leafsearch là hệ thống Content-Based Image Retrieval (CBIR) cho ảnh lá. "
        "Từ một ảnh lá đầu vào, pipeline trích xuất bốn nhóm đặc trưng riêng biệt — "
        "hình dạng, màu sắc, kết cấu và gân lá — tổng cộng 125 chiều, lưu vào "
        "PostgreSQL với pgvector, và tìm kiếm bằng Cosine distance có trọng số.",
    )

    add_heading(doc, "1.1. Cấu trúc vector", level=2)
    add_dim_table(
        doc,
        headers=["Nhóm", "Số chiều", "Cột DB", "Nội dung tóm tắt"],
        rows=[
            ["Shape",    "10", "shape_vector",    "7 đặc trưng hình học + 3 Hu moments log-transformed"],
            ["Color",    "72", "color_vector",    "Radial Spatial HSV: 3 vòng đồng tâm × 24 bin (8H×3S)"],
            ["Texture",  "38", "texture_vector",  "4 GLCM Haralick + 24 Gabor (12 filter × mean+var) + 10 LBP riu2"],
            ["Venation", "5",  "venation_vector", "Density, edge density, thickness, contrast, uniformity"],
            ["TỔNG",     "125", "—",              "—"],
        ],
    )

    add_heading(doc, "1.2. Luồng xử lý", level=2)
    add_bullet(doc, "Offline — build dataset: ingest.py → normalize.py (Z-score per-group).")
    add_bullet(doc, "Online — query: upload → preprocess → extract 4 vector → Z-score → 4 cosine + weighted sum → top-K.")

    add_para(doc, "Trình tự gọi mã nguồn (online query):", italic=True)
    add_code(doc,
        "  app.py\n"
        "    └─ features.preprocess.preprocess()            # tách nền + resize 256×256\n"
        "    └─ features.pipeline.extract_all()             # trả dict 4 vector\n"
        "    └─ features.zscore.normalize_all()             # Z-score per-group\n"
        "    └─ db.crud.search_weighted_cosine()            # 4 cosine + weighted sum"
    )


def section_preprocess(doc):
    add_heading(doc, "2. Tiền xử lý ảnh", level=1)

    add_para(doc,
        "Mọi ảnh vào pipeline đều phải qua hai bước tiền xử lý cố định, được implement trong "
        "features/preprocess.py. Pipeline cũ có thêm bước xoay lá theo trục dọc (PCA-based) "
        "nhưng đã bị loại bỏ vì gây sai lệch trên các lá tròn/lá có khuyết, và mọi extractor "
        "trong hệ thống đều rotation-invariant (hoặc tự xử lý orientation nội bộ).",
    )

    add_heading(doc, "2.1. Tách nền (segment_leaf)", level=2)
    add_kv_table(doc, [
        ("Mục tiêu",  "Tạo mặt nạ nhị phân của lá trên nền đen, loại bỏ noise nhỏ."),
        ("Bước 1",    "Chuyển grayscale, làm mờ Gaussian kernel 7×7."),
        ("Bước 2",    "Otsu threshold (THRESH_BINARY_INV + THRESH_OTSU) — auto chọn ngưỡng tối ưu giữa nền sáng và lá tối hơn."),
        ("Bước 3",    "Morphological close với kernel ellipse 5×5 — vá lỗ nhỏ trong mask."),
        ("Bước 4",    "_keep_largest_component() — giữ contour có diện tích lớn nhất, loại bỏ blob lẻ."),
        ("Output",    "leaf_image (BGR với nền đen) + leaf_mask (uint8 nhị phân, lá=255)."),
    ])

    add_heading(doc, "2.2. Resize + pad (resize_and_pad)", level=2)
    add_kv_table(doc, [
        ("Mục tiêu",  "Đưa mọi ảnh về kích thước cố định 256×256 mà không bóp méo tỷ lệ."),
        ("Công thức", "ratio = target_size / max(h, w)\nnew_h, new_w = round(h·ratio), round(w·ratio)\nResize với cv2.INTER_AREA"),
        ("Pad",       "BORDER_CONSTANT giá trị [0,0,0] (đen) — pad đều 2 bên để giữ lá ở giữa."),
        ("TARGET_SIZE", "256 (hằng số trong features/preprocess.py)"),
    ])

    add_heading(doc, "2.3. Hàm hỗ trợ Unicode đường dẫn", level=2)
    add_para(doc,
        "Trên Windows, cv2.imread() không hỗ trợ đường dẫn chứa ký tự Unicode tiếng Việt. "
        "Hệ thống dùng numpy.fromfile() + cv2.imdecode() để đảm bảo đọc được mọi đường dẫn. "
        "Triển khai trong features/preprocess.py:preprocess() và features/extractors/base.py:_read_image()."
    )


def section_shape(doc):
    add_heading(doc, "3. Nhóm 1: Đặc trưng hình dạng (Shape) — 10 chiều", level=1)
    add_para(doc,
        "File: features/extractors/shape.py — class ShapeExtractor. "
        "Đầu vào: ảnh BGR đã preprocess; mask được tính từ BaseExtractor._get_leaf_mask() "
        "(threshold=10 trên grayscale, giả định nền đen). Toàn bộ đặc trưng tính trên contour lớn nhất.",
        italic=True,
    )

    add_heading(doc, "3.1. Bảng tóm tắt 10 chiều", level=2)
    add_dim_table(
        doc,
        headers=["#", "Tên đặc trưng", "Công thức", "Khoảng giá trị", "Ý nghĩa"],
        rows=[
            ["F01", "aspect_ratio", "min(w,h) / max(w,h)", "(0, 1]", "Tỷ lệ khung lá: gần 0 = lá dài hẹp, 1 = tròn/vuông"],
            ["F02", "solidity", "area / convex_area", "(0, 1]", "Độ đặc — gần 1 = mép nhẵn, thấp = lá xẻ thùy"],
            ["F03", "circularity", "4πA / P²", "(0, 1]", "Độ tròn — 1 = tròn hoàn hảo"],
            ["F04", "convexity", "perimeter_hull / perimeter_contour", "(0, 1]", "Gần 1 = mép trơn, thấp = nhiều răng cưa"],
            ["F05", "extent", "area / (w·h)", "(0, 1]", "Tỷ lệ lấp đầy bounding box"],
            ["F06", "eccentricity", "√(1 − b²/a²) từ fitEllipse", "[0, 1)", "0 = tròn, gần 1 = ellipse dẹt"],
            ["F07", "relative_center_of_mass", "min(cy − y_bb, y_bb + h − cy) / h", "[0, 0.5]", "Vị trí trọng tâm so với mép trên/dưới gần hơn"],
            ["F08", "hu_moment_1", "−sign(h₁)·log₁₀|h₁|", "≈ [0, 20]", "Hu moment bậc 1, log-transformed"],
            ["F09", "hu_moment_2", "−sign(h₂)·log₁₀|h₂|", "≈ [0, 20]", "Hu moment bậc 2, log-transformed"],
            ["F10", "hu_moment_3", "−sign(h₃)·log₁₀|h₃|", "≈ [0, 20]", "Hu moment bậc 3, log-transformed"],
        ],
    )

    add_heading(doc, "3.2. Chi tiết tính toán", level=2)
    add_para(doc, "Pseudocode (rút gọn từ shape.py):", italic=True)
    add_code(doc,
        "mask = _get_leaf_mask(image)                                  # threshold(10) trên grayscale\n"
        "contours, _ = cv2.findContours(mask, RETR_EXTERNAL, ...)\n"
        "c = max(contours, key=cv2.contourArea)\n"
        "\n"
        "area      = cv2.contourArea(c)\n"
        "perimeter = cv2.arcLength(c, True)\n"
        "x, y, w, h = cv2.boundingRect(c)\n"
        "hull = cv2.convexHull(c)\n"
        "hull_area, hull_perimeter = cv2.contourArea(hull), cv2.arcLength(hull, True)\n"
        "\n"
        "# F06 eccentricity (yêu cầu ≥ 5 điểm contour)\n"
        "(_, _), (MA, ma), _ = cv2.fitEllipse(c)\n"
        "a, b = max(MA,ma)/2, min(MA,ma)/2\n"
        "ecc  = sqrt(1 − b²/a²)\n"
        "\n"
        "# F07 RCM — khoảng cách từ trọng tâm tới mép gần hơn / chiều cao bbox\n"
        "moments = cv2.moments(c)\n"
        "cy = moments['m01'] / moments['m00']\n"
        "rcm = min(cy − y, (y+h) − cy) / h\n"
        "\n"
        "# F08-F10 Hu moments với log-transform\n"
        "hu_raw = cv2.HuMoments(moments).flatten()\n"
        "hu_log = [−sign(v) · log10(|v| + 1e-10) for v in hu_raw[:3]]"
    )

    add_para(doc,
        "Các đặc trưng F02, F03, F04, F06 và Hu moments đều có tính bất biến với phép xoay, "
        "phóng/thu, và tịnh tiến. F01, F05, F07 phụ thuộc hướng (orientation-sensitive) nhưng vẫn "
        "ổn định trên dataset có lá ảnh đứng tự nhiên.",
    )


def section_color(doc):
    add_heading(doc, "4. Nhóm 2: Đặc trưng màu sắc (Color) — 72 chiều", level=1)
    add_para(doc,
        "File: features/extractors/color.py — class ColorExtractor. "
        "Đặc trưng có dạng Radial Spatial Histogram HSV: chia vùng lá thành ba vòng đồng tâm "
        "tính từ trọng tâm ra mép, mỗi vòng tính một histogram 2D Hue × Saturation.",
        italic=True,
    )

    add_heading(doc, "4.1. Lượng tử hóa và bố cục vector", level=2)
    add_kv_table(doc, [
        ("Kênh Hue (H)",       "8 bin trải đều [0, 180) — OpenCV scale"),
        ("Kênh Saturation (S)", "3 bin trải đều [0, 256)"),
        ("Kênh Value (V)",      "KHÔNG dùng — tránh nhiễu do bóng/ánh sáng"),
        ("Bin/vòng",           "8 × 3 = 24 bin"),
        ("Số vòng",            "3 (Core / Middle / Periphery)"),
        ("Tổng chiều",         "24 × 3 = 72"),
        ("Sắp xếp",            "[ring1 24 bin | ring2 24 bin | ring3 24 bin]"),
        ("Index 2D → 1D",      "hist2d.flatten()  — bin (h_idx, s_idx) → h_idx · S_BINS + s_idx"),
    ])

    add_heading(doc, "4.2. Tính 3 vòng đồng tâm", level=2)
    add_para(doc, "Bước 1 — Trọng tâm lá:", bold=True)
    add_code(doc,
        "M = cv2.moments(mask)\n"
        "cx, cy = M['m10']/M['m00'], M['m01']/M['m00']"
    )
    add_para(doc, "Bước 2 — Bản đồ khoảng cách đến trọng tâm:", bold=True)
    add_code(doc,
        "ys, xs = np.indices(mask.shape)\n"
        "dist = sqrt((xs − cx)² + (ys − cy)²)"
    )
    add_para(doc, "Bước 3 — Ngưỡng 3 vòng:", bold=True)
    add_code(doc,
        "R_max = dist[mask > 0].max()\n"
        "ring1: 0 ≤ d < R_max/3       (Core — vùng lõi gần trọng tâm)\n"
        "ring2: R_max/3 ≤ d < 2R_max/3 (Middle — vùng giữa)\n"
        "ring3: 2R_max/3 ≤ d ≤ R_max   (Periphery — vùng rìa, gần mép)"
    )

    add_heading(doc, "4.3. Histogram 2D H×S trong từng vòng", level=2)
    add_code(doc,
        "for r_idx, ring_mask in enumerate([m1, m2, m3]):\n"
        "    H_pixels = hsv[:, :, 0][ring_mask]\n"
        "    S_pixels = hsv[:, :, 1][ring_mask]\n"
        "    hist2d, _, _ = np.histogram2d(\n"
        "        H_pixels, S_pixels,\n"
        "        bins=[8, 3],\n"
        "        range=[[0, 180], [0, 256]],\n"
        "    )\n"
        "    hist2d /= (hist2d.sum() + 1e-10)   # chuẩn hóa theo tổng pixel của vòng\n"
        "    ring_vector = hist2d.flatten()       # 24 chiều"
    )

    add_heading(doc, "4.4. Tính bất biến và lợi ích", level=2)
    add_bullet(doc, "Bất biến xoay: khoảng cách pixel → trọng tâm không đổi khi xoay ảnh.")
    add_bullet(doc, "Bất biến tỷ lệ: chuẩn hóa từng vòng theo tổng pixel của vòng đó.")
    add_bullet(doc, "Spatial-aware: tách màu vùng lõi (đậm hơn) khỏi rìa (thường nhạt/khô) → phân biệt loài tốt hơn so với histogram phẳng.")

    add_para(doc,
        "Lưu ý thực tế: vì lá phần lớn chỉ có hue xanh (≈ 30–80 trong scale OpenCV) và phần nền "
        "có saturation thấp, nhiều bin của histogram luôn = 0 trên cả dataset. Trong bước Z-score, "
        "22 chiều của 72 chiều bị cảnh báo std ≈ 0 — chúng sẽ collapse về 0 sau normalize và không "
        "đóng góp vào cosine distance.",
        italic=True,
    )


def section_texture(doc):
    add_heading(doc, "5. Nhóm 3: Đặc trưng kết cấu (Texture) — 38 chiều", level=1)
    add_para(doc,
        "File: features/extractors/texture.py — class TextureExtractor. "
        "Ba khối phụ: GLCM Haralick (4 chiều) + Gabor bank (24 chiều) + LBP riu2 (10 chiều).",
        italic=True,
    )

    add_heading(doc, "5.1. GLCM Haralick — chiều [0:4]", level=2)
    add_kv_table(doc, [
        ("Thư viện",  "skimage.feature.graycomatrix, graycoprops"),
        ("Distances", "[1] — 1 pixel"),
        ("Angles",    "[0°, 45°, 90°, 135°] — bất biến xoay nhờ lấy mean"),
        ("Levels",    "256 — full grayscale, không quantize"),
        ("Symmetric", "True"),
        ("Normed",    "True"),
    ])
    add_dim_table(
        doc,
        headers=["#", "Tên", "Công thức", "Ý nghĩa"],
        rows=[
            ["F83", "glcm_contrast",    "Σ(i−j)² · P(i,j)",                  "Tương phản — cao khi bề mặt gồ ghề, gân nổi rõ"],
            ["F84", "glcm_energy",      "Σ P(i,j)²",                          "Năng lượng — cao khi bề mặt nhẵn, phân bố đồng đều"],
            ["F85", "glcm_homogeneity", "Σ P(i,j) / (1 + (i−j)²)",            "Độ đồng nhất — cao khi biểu bì lá mượt mà"],
            ["F86", "glcm_correlation", "Đo phụ thuộc tuyến tính pixel lân cận", "Cao khi có cấu trúc có hướng (gân song song)"],
        ],
    )
    add_para(doc, "Mỗi giá trị là trung bình của 4 phép tính theo 4 góc → bất biến xoay.", italic=True)

    add_heading(doc, "5.2. Gabor filter bank — chiều [4:28]", level=2)
    add_kv_table(doc, [
        ("Số hướng",     "4 — θ ∈ {0°, 45°, 90°, 135°}"),
        ("Số scale",     "3 — (σ, λ) ∈ {(2,4), (4,8), (6,12)}"),
        ("Tổng filter",  "4 × 3 = 12"),
        ("Kernel size",  "21 × 21"),
        ("Gamma",        "0.5"),
        ("Psi",          "0.0"),
        ("API",          "cv2.getGaborKernel + cv2.filter2D"),
        ("Chiều/filter", "2 — (mean, variance) của response trên vùng lá"),
        ("Tổng chiều",   "12 × 2 = 24"),
    ])
    add_para(doc, "Bố cục sắp xếp:", italic=True)
    add_code(doc,
        "[θ0_s0_mean, θ0_s0_var, θ0_s1_mean, θ0_s1_var, θ0_s2_mean, θ0_s2_var,\n"
        " θ1_s0_mean, θ1_s0_var, ...                                       ,\n"
        " ...                                                              ,\n"
        " θ3_s2_mean, θ3_s2_var]                       # 12 filter × 2 = 24"
    )
    add_para(doc, "Tính toán cho mỗi filter:", italic=True)
    add_code(doc,
        "kernel   = cv2.getGaborKernel((21,21), sigma=σ, theta=θ, lambd=λ, gamma=0.5, psi=0)\n"
        "filtered = cv2.filter2D(gray_f32, cv2.CV_32F, kernel)\n"
        "vals     = filtered[mask > 0]\n"
        "feat_mean = np.mean(vals)\n"
        "feat_var  = np.var(vals)"
    )

    add_heading(doc, "5.3. LBP riu2 histogram — chiều [28:38]", level=2)
    add_kv_table(doc, [
        ("Phương pháp", "Rotation Invariant Uniform LBP (riu2)"),
        ("Thư viện",    "skimage.feature.local_binary_pattern(method='uniform')"),
        ("P",           "8 (số điểm lân cận)"),
        ("R",           "1 (bán kính)"),
        ("Số bin",      "10 — P + 2"),
        ("Bin 0–8",     "Uniform patterns có 0–8 bit '1'"),
        ("Bin 9",       "Tất cả non-uniform patterns (gom chung)"),
        ("Chuẩn hóa",   "hist / hist.sum() — tỷ lệ pixel trong vùng lá"),
    ])
    add_code(doc,
        "lbp = local_binary_pattern(gray, P=8, R=1, method='uniform')\n"
        "vals = lbp[mask > 0]\n"
        "hist, _ = np.histogram(vals, bins=np.arange(11))   # 10 bin\n"
        "hist = hist / (hist.sum() + 1e-10)"
    )
    add_para(doc, "LBP riu2 bất biến với xoay vì các pattern được nhóm theo số bit '1' "
                  "thay vì vị trí cụ thể của các bit.", italic=True)


def section_venation(doc):
    add_heading(doc, "6. Nhóm 4: Đặc trưng gân lá (Venation) — 5 chiều", level=1)
    add_para(doc,
        "File: features/extractors/vein_extractor.py — class VeinExtractor. "
        "Toàn bộ đặc trưng tính trên mặt nạ gân (veins_mask) chiết xuất bằng Blackhat morphology.",
        italic=True,
    )

    add_heading(doc, "6.1. Pipeline tiền xử lý chung", level=2)
    add_code(doc,
        "leaf_mask = _get_leaf_mask(image)              # threshold(10) trên grayscale\n"
        "gray      = cv2.cvtColor(image, COLOR_BGR2GRAY)\n"
        "\n"
        "# Blackhat làm nổi gân (vùng tối) trên nền lá (vùng sáng)\n"
        "kernel = cv2.getStructuringElement(MORPH_ELLIPSE, (15, 15))\n"
        "blackhat = cv2.morphologyEx(gray, MORPH_BLACKHAT, kernel)\n"
        "\n"
        "# Threshold → veins_mask\n"
        "_, veins_mask = cv2.threshold(blackhat, 15, 255, THRESH_BINARY)\n"
        "veins_inside  = bitwise_and(veins_mask, leaf_mask)\n"
        "\n"
        "# Canny trên gân để bắt cạnh tinh\n"
        "edges_mask    = cv2.Canny(veins_inside, 50, 150)\n"
        "edges_inside  = bitwise_and(edges_mask, leaf_mask)"
    )

    add_heading(doc, "6.2. Bảng 5 đặc trưng", level=2)
    add_dim_table(
        doc,
        headers=["#", "Tên", "Công thức", "Đơn vị", "Ý nghĩa"],
        rows=[
            ["F121", "vein_density",      "vein_area / leaf_area",                "[0, 1]",  "Mật độ gân — cao = gân dày đặc"],
            ["F122", "vein_edge_density", "edge_area / leaf_area",                "[0, 1]",  "Mật độ cạnh gân — cao = mạng gân phức tạp"],
            ["F123", "vein_thickness",    "vein_area / (edge_area + 1e-6)",       "≥ 0",     "Độ dày sợi gân ước lượng"],
            ["F124", "vein_contrast",     "mean(blackhat[veins_inside > 0])",      "[0, 255]", "Cường độ blackhat trung bình — gân nổi bật hay mờ"],
            ["F125", "vein_uniformity",   "std(blackhat[veins_inside > 0])",       "[0, ~128]", "Độ lệch chuẩn — thấp = gân đều, cao = gân đậm nhạt xen kẽ"],
        ],
    )

    add_heading(doc, "6.3. Phân biệt loài qua đặc trưng gân", level=2)
    add_bullet(doc, "Lá ổi/táo: vein_density cao, vein_edge_density cao (mạng gân chằng chịt).")
    add_bullet(doc, "Lá hành/tre: vein_density và vein_edge_density thấp, vein_uniformity thấp (gân song song đều).")
    add_bullet(doc, "Lá bàng/bồ đề: vein_thickness cao (gân chính to thô); vein_contrast cao.")


def section_normalize(doc):
    add_heading(doc, "7. Chuẩn hóa Z-score per-group", level=1)
    add_para(doc,
        "File: features/zscore.py + scripts/normalize.py. "
        "Z-score được áp dụng RIÊNG cho từng nhóm vector — không trộn chung 4 nhóm — vì "
        "scale của Gabor variance (đơn vị bình phương cường độ pixel) khác hẳn so với histogram "
        "HSV (xác suất [0,1]).",
        italic=True,
    )

    add_heading(doc, "7.1. Quy trình offline (scripts/normalize.py)", level=2)
    add_code(doc,
        "for group in ['shape', 'color', 'texture', 'venation']:\n"
        "    X = ma trận N × D_group  (N = 1907 ảnh)\n"
        "    μ = X.mean(axis=0)\n"
        "    σ = X.std(axis=0)\n"
        "    X_norm = (X − μ) / (σ + 1e-8)\n"
        "    overwrite cột tương ứng trong DB"
    )
    add_para(doc, "Output: data/zscore_params.npz chứa 8 mảng:")
    add_code(doc, "shape_mean, shape_std, color_mean, color_std,\n"
                  "texture_mean, texture_std, venation_mean, venation_std")

    add_heading(doc, "7.2. Quy trình online (features/zscore.py)", level=2)
    add_code(doc,
        "from features.zscore import normalize_all\n"
        "from features.pipeline import extract_all\n"
        "\n"
        "raw  = extract_all('uploaded.jpg')      # {'shape':[..10], 'color':[..72], ...}\n"
        "norm = normalize_all(raw)               # áp dụng (μ, σ) per-group, lazy-load từ npz"
    )
    add_para(doc,
        "Điểm quan trọng: query và DB phải dùng CÙNG (μ, σ). Nếu chạy lại normalize.py "
        "thì zscore_params.npz mới được ghi và features.zscore tự đọc lại lần gọi tiếp theo.",
        italic=True,
    )


def section_db_schema(doc):
    add_heading(doc, "8. Lưu trữ database", level=1)
    add_para(doc, "File: db/models.py + db/database.py. PostgreSQL 16 + pgvector extension.", italic=True)

    add_heading(doc, "8.1. Schema bảng leaf_images", level=2)
    add_code(doc,
        "CREATE TABLE leaf_images (\n"
        "    image_id        INTEGER PRIMARY KEY,\n"
        "    file_name       VARCHAR NOT NULL,\n"
        "    file_path       VARCHAR NOT NULL,\n"
        "    shape_vector    VECTOR(10),\n"
        "    color_vector    VECTOR(72),\n"
        "    texture_vector  VECTOR(38),\n"
        "    venation_vector VECTOR(5)\n"
        ");"
    )

    add_heading(doc, "8.2. Quy ước image_id", level=2)
    add_bullet(doc, "image_id = int(filename_stem) — ví dụ 1001.jpg → 1001.")
    add_bullet(doc, "scripts/ingest.py giả định mọi filename trong data/processed là số nguyên.")
    add_bullet(doc, "Tái chạy ingest.py với --force để overwrite vector mà không tạo trùng.")


def section_search(doc):
    add_heading(doc, "9. Tìm kiếm tương đồng", level=1)
    add_para(doc, "File: db/crud.py — hàm search_weighted_cosine.", italic=True)

    add_heading(doc, "9.1. Công thức tổ hợp", level=2)
    add_code(doc,
        "Total = W_shape    · (shape_vector    <=> q_shape)\n"
        "      + W_color    · (color_vector    <=> q_color)\n"
        "      + W_texture  · (texture_vector  <=> q_texture)\n"
        "      + W_venation · (venation_vector <=> q_venation)\n"
        "\n"
        "ORDER BY Total ASC  LIMIT K"
    )
    add_para(doc, "<=> là toán tử cosine distance của pgvector. Trọng số mặc định 0.25 đều nhau.")

    add_heading(doc, "9.2. Trọng số trong UI Streamlit", level=2)
    add_bullet(doc, "Sidebar có 4 slider (Shape / Color / Texture / Venation), giá trị 0.0–1.0.")
    add_bullet(doc, "Sau khi nhập, hệ thống tự chuẩn hóa tổng = 1 trước khi gửi xuống DB.")
    add_bullet(doc, "Kết quả hiển thị kèm breakdown từng thành phần distance để debug.")

    add_heading(doc, "9.3. Truy vấn SQL thực tế", level=2)
    add_code(doc,
        "SELECT image_id, file_name,\n"
        "       (shape_vector    <=> CAST(:v_shape    AS vector)) AS d_shape,\n"
        "       (color_vector    <=> CAST(:v_color    AS vector)) AS d_color,\n"
        "       (texture_vector  <=> CAST(:v_texture  AS vector)) AS d_texture,\n"
        "       (venation_vector <=> CAST(:v_venation AS vector)) AS d_venation,\n"
        "       (w_shape*d_shape + w_color*d_color\n"
        "        + w_texture*d_texture + w_venation*d_venation) AS total\n"
        "FROM leaf_images\n"
        "ORDER BY total ASC\n"
        "LIMIT :k"
    )


def section_summary(doc):
    add_heading(doc, "10. Tổng kết", level=1)
    add_para(doc,
        "Hệ thống Leafsearch sử dụng đa đặc trưng (multi-feature) thay vì một deep embedding "
        "duy nhất. Mỗi nhóm đặc trưng nắm bắt một khía cạnh khác nhau của lá — hình dạng tổng "
        "thể, phân bố màu theo vị trí, kết cấu bề mặt, và mạng gân — và tổ hợp chúng linh hoạt "
        "thông qua trọng số trong cosine search."
    )

    add_heading(doc, "10.1. Ưu điểm hệ thống hiện tại", level=2)
    add_bullet(doc, "Toàn bộ trích xuất bằng OpenCV + scikit-image, không cần GPU/PyTorch.")
    add_bullet(doc, "Đặc trưng rotation-invariant ở phần lớn các chiều (Hu, GLCM mean, Gabor, LBP riu2, radial histogram).")
    add_bullet(doc, "Có thể giải thích từng chiều — phù hợp đồ án nghiên cứu, dễ debug.")
    add_bullet(doc, "Tìm kiếm < 50 ms cho 1907 ảnh nhờ pgvector + cosine native operator.")
    add_bullet(doc, "Trọng số tổ hợp tùy chỉnh thời gian thực qua sidebar Streamlit.")

    add_heading(doc, "10.2. Hạn chế cần lưu ý", level=2)
    add_bullet(doc, "22/72 chiều của color_vector có std ≈ 0 — phần dataset hiện tại đồng đều về hue → một số bin không bao giờ kích hoạt.")
    add_bullet(doc, "GLCM 256 levels chính xác nhưng tốn nhiều bộ nhớ (~250 KB × 4 ma trận / ảnh).")
    add_bullet(doc, "Pipeline không có deep learning embedding — không cạnh tranh với mô hình CNN trên dataset rất phức tạp.")

    add_heading(doc, "10.3. Hướng mở rộng", level=2)
    add_bullet(doc, "Thêm CNN embedding (ResNet/EfficientNet) làm nhóm thứ 5 để bắt đặc trưng ngữ nghĩa cao.")
    add_bullet(doc, "Tạo HNSW index per-cột để tăng tốc truy vấn khi dataset > 100k ảnh.")
    add_bullet(doc, "Cân nhắc bỏ các chiều color std≈0 hoặc dùng L1-norm thay vì histogram thường.")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Page margin gọn
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section_cover(doc)
    section_overview(doc)
    section_preprocess(doc)
    section_shape(doc)
    section_color(doc)
    section_texture(doc)
    section_venation(doc)
    section_normalize(doc)
    section_db_schema(doc)
    section_search(doc)
    section_summary(doc)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
